"""
CV Profiler edge-case tests (Phase 7 / QA — weird PDFs).

Covers the failure modes the parser is expected to surface cleanly:

    * Zero-byte upload.
    * File over the 10MB limit.
    * Image-only PDF (valid structure, but no selectable text).
    * Corrupted PDF header.
    * Unknown / spoofed extension.
    * DOCX with no readable paragraphs.
    * Valid PDF / DOCX happy paths — sanity anchor for the suite.

We construct every fixture on the fly with ``pypdf`` + ``python-docx`` so
nothing binary enters the repo.
"""

from __future__ import annotations

import io
import os
from pathlib import Path

import pypdf
import pytest
from docx import Document

from app.utils.file_processing import (
    MAX_FILE_SIZE,
    extract_text_from_file,
    validate_cv_file,
)

# ----------------------------------------------------------------------
# Fixture builders (pure functions on a tmp_path)
# ----------------------------------------------------------------------


def _make_valid_pdf(tmp_path: Path, text: str = "Hello") -> Path:
    """Build a minimal one-page PDF with ``text`` on it."""
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    target = tmp_path / "valid.pdf"
    with open(target, "wb") as handle:
        writer.write(handle)
    # pypdf.add_blank_page yields a valid but text-less PDF; that's OK for
    # validation. Content-text tests use _make_pdf_with_content below.
    return target


def _make_pdf_with_content(tmp_path: Path, body: str) -> Path:
    """Build a real PDF containing extractable ``body`` text via reportlab."""
    from reportlab.pdfgen import canvas

    target = tmp_path / "content.pdf"
    c = canvas.Canvas(str(target))
    c.drawString(72, 720, body)
    c.showPage()
    c.save()
    return target


def _make_empty_file(tmp_path: Path, name: str) -> Path:
    """Zero-byte file with the given name."""
    target = tmp_path / name
    target.touch()
    return target


def _make_oversized_pdf(tmp_path: Path) -> Path:
    """Valid PDF that sits just above the MAX_FILE_SIZE limit."""
    path = _make_valid_pdf(tmp_path)
    # Append garbage bytes to push the file over the 10 MB ceiling.
    padding = MAX_FILE_SIZE + 1024
    with open(path, "ab") as handle:
        handle.write(b"\0" * padding)
    return path


def _make_corrupted_pdf(tmp_path: Path) -> Path:
    """File with a ``.pdf`` extension but non-PDF content."""
    target = tmp_path / "corrupt.pdf"
    target.write_bytes(b"not a real pdf \x00 just garbage")
    return target


def _make_valid_docx(tmp_path: Path, body: str) -> Path:
    target = tmp_path / "valid.docx"
    doc = Document()
    doc.add_paragraph(body)
    doc.save(str(target))
    return target


def _make_empty_docx(tmp_path: Path) -> Path:
    """Technically-valid DOCX with no paragraphs."""
    target = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(target))
    return target


# ----------------------------------------------------------------------
# validate_cv_file
# ----------------------------------------------------------------------


class TestValidateCvFile:
    def test_accepts_valid_pdf(self, tmp_path: Path) -> None:
        path = _make_valid_pdf(tmp_path)
        ok, err = validate_cv_file(str(path), "valid.pdf")
        assert ok, err

    def test_accepts_valid_docx(self, tmp_path: Path) -> None:
        path = _make_valid_docx(tmp_path, "John Doe — Software Engineer")
        ok, err = validate_cv_file(str(path), "valid.docx")
        assert ok, err

    def test_rejects_nonexistent_path(self, tmp_path: Path) -> None:
        ok, err = validate_cv_file(str(tmp_path / "missing.pdf"), "missing.pdf")
        assert not ok
        assert "does not exist" in err.lower()

    def test_rejects_zero_byte_file(self, tmp_path: Path) -> None:
        path = _make_empty_file(tmp_path, "blank.pdf")
        ok, err = validate_cv_file(str(path), "blank.pdf")
        assert not ok
        assert "empty" in err.lower()

    def test_rejects_oversized_file(self, tmp_path: Path) -> None:
        path = _make_oversized_pdf(tmp_path)
        # Sanity: the fixture must actually be over the limit.
        assert os.path.getsize(path) > MAX_FILE_SIZE
        ok, err = validate_cv_file(str(path), "huge.pdf")
        assert not ok
        assert "exceeds limit" in err.lower()

    @pytest.mark.parametrize(
        "filename",
        ["resume.txt", "resume.rtf", "resume.doc", "resume.exe"],
    )
    def test_rejects_unsupported_extensions(
        self, tmp_path: Path, filename: str
    ) -> None:
        path = _make_valid_pdf(tmp_path)
        ok, err = validate_cv_file(str(path), filename)
        assert not ok
        # Either the extension check fires first, or the MIME check fires.
        assert "not supported" in err.lower() or "not recognized" in err.lower()

    def test_rejects_corrupted_pdf(self, tmp_path: Path) -> None:
        path = _make_corrupted_pdf(tmp_path)
        ok, err = validate_cv_file(str(path), "corrupt.pdf")
        assert not ok
        assert "invalid pdf" in err.lower() or "pdf" in err.lower()


# ----------------------------------------------------------------------
# extract_text_from_file
# ----------------------------------------------------------------------


class TestExtractTextFromFile:
    def test_extracts_text_from_real_pdf(self, tmp_path: Path) -> None:
        path = _make_pdf_with_content(tmp_path, "Software Engineer since 2019")
        text = extract_text_from_file(str(path), "real.pdf")
        assert "Software Engineer" in text

    def test_extracts_text_from_docx(self, tmp_path: Path) -> None:
        path = _make_valid_docx(tmp_path, "Backend dev with FastAPI and Postgres")
        text = extract_text_from_file(str(path), "cv.docx")
        assert "FastAPI" in text
        assert "Postgres" in text

    def test_image_only_pdf_raises_no_text(self, tmp_path: Path) -> None:
        """PDF with no extractable text (e.g. scanned image)."""
        # _make_valid_pdf produces a blank page — no extractable strings.
        path = _make_valid_pdf(tmp_path)
        with pytest.raises(ValueError, match="no extractable text|no pages"):
            extract_text_from_file(str(path), "scan.pdf")

    def test_empty_docx_raises(self, tmp_path: Path) -> None:
        path = _make_empty_docx(tmp_path)
        with pytest.raises((ValueError, Exception)):
            extract_text_from_file(str(path), "empty.docx")

    def test_corrupted_pdf_raises(self, tmp_path: Path) -> None:
        path = _make_corrupted_pdf(tmp_path)
        with pytest.raises((ValueError, Exception)):
            extract_text_from_file(str(path), "corrupt.pdf")

    def test_unsupported_extension_raises(self, tmp_path: Path) -> None:
        # Even if content is PDF, a .txt name must be rejected.
        path = _make_pdf_with_content(tmp_path, "hello")
        renamed = tmp_path / "cv.txt"
        os.rename(path, renamed)
        with pytest.raises(ValueError):
            extract_text_from_file(str(renamed), "cv.txt")
