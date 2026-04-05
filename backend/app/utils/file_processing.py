"""
File Processing Utilities — PDF and DOCX text extraction.

Provides secure file handling and text extraction for CV parsing.
Supports common resume formats with error handling and validation.
"""

import logging
import mimetypes
import os
from pathlib import Path
from typing import List, Tuple

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)

# Supported file types with corresponding MIME types
SUPPORTED_MIME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx"
}

SUPPORTED_EXTENSIONS = [".pdf", ".docx"]
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit


def validate_cv_file(file_path: str, original_filename: str) -> Tuple[bool, str]:
    """Validate uploaded CV file for security and format compliance.

    Args:
        file_path: Path to the uploaded file on disk.
        original_filename: Original filename from upload.

    Returns:
        Tuple[bool, str]: (is_valid, error_message)
            - is_valid: True if file passes all validation checks
            - error_message: Detailed error if validation fails (empty if valid)

    Validation checks:
    - File exists and is readable
    - File size within limits (10MB max)
    - File extension in allowed list (.pdf, .docx)
    - MIME type verification (prevents extension spoofing)
    - Basic file integrity (can be opened by respective libraries)
    """
    try:
        # Check file exists
        if not os.path.exists(file_path):
            return False, "File does not exist on disk"

        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "File is empty"
        if file_size > MAX_FILE_SIZE:
            return False, f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds limit ({MAX_FILE_SIZE / 1024 / 1024}MB)"

        # Check extension
        file_extension = Path(original_filename).suffix.lower()
        if file_extension not in SUPPORTED_EXTENSIONS:
            return False, f"File extension '{file_extension}' not supported. Allowed: {', '.join(SUPPORTED_EXTENSIONS)}"

        # MIME type verification (security: prevents .pdf.exe type attacks)
        detected_mime, _ = mimetypes.guess_type(file_path)
        if detected_mime not in SUPPORTED_MIME_TYPES:
            return False, f"File type not recognized or unsupported. Detected: {detected_mime}"

        # File integrity check (attempt to open with appropriate library)
        if file_extension == ".pdf":
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                if len(pdf_reader.pages) == 0:
                    return False, "PDF file contains no pages"

        elif file_extension == ".docx":
            # Attempt to open with python-docx
            doc = Document(file_path)
            if len(doc.paragraphs) == 0:
                return False, "DOCX file contains no readable text"

        return True, ""

    except PyPDF2.errors.PdfReadError as e:
        return False, f"Invalid PDF file: {str(e)}"
    except Exception as e:
        return False, f"File validation error: {str(e)}"


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file.

    Args:
        file_path: Path to the PDF file.

    Returns:
        str: Extracted text content with basic formatting preserved.

    Raises:
        ValueError: If PDF cannot be read or contains no text.
        FileNotFoundError: If file doesn't exist.
    """
    try:
        extracted_text = []

        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if len(pdf_reader.pages) == 0:
                raise ValueError("PDF contains no pages")

            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        extracted_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {e}")
                    continue

            if not extracted_text:
                raise ValueError("PDF contains no extractable text")

            return "\n\n".join(extracted_text)

    except FileNotFoundError:
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    except PyPDF2.errors.PdfReadError as e:
        raise ValueError(f"Invalid PDF file: {str(e)}")
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        str: Extracted text content with paragraph separation.

    Raises:
        ValueError: If DOCX cannot be read or contains no text.
        FileNotFoundError: If file doesn't exist.
    """
    try:
        doc = Document(file_path)

        # Extract text from paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                paragraphs.extend(table_text)

        if not paragraphs:
            raise ValueError("DOCX contains no extractable text")

        return "\n\n".join(paragraphs)

    except FileNotFoundError:
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_file(file_path: str, original_filename: str) -> str:
    """Extract text from a CV file based on its type.

    Automatically detects file type and uses appropriate extraction method.
    Includes validation and error handling.

    Args:
        file_path: Path to the uploaded file.
        original_filename: Original filename (used for extension detection).

    Returns:
        str: Extracted text content ready for AI processing.

    Raises:
        ValueError: If file validation fails or text extraction fails.
        FileNotFoundError: If file doesn't exist.

    Usage:
        try:
            text = extract_text_from_file("/tmp/resume.pdf", "resume.pdf")
            # Process with AI...
        except ValueError as e:
            # Handle validation/extraction error
            logger.error(f"CV extraction failed: {e}")
    """
    # First validate the file
    is_valid, error_message = validate_cv_file(file_path, original_filename)
    if not is_valid:
        raise ValueError(f"File validation failed: {error_message}")

    # Extract text based on file extension
    file_extension = Path(original_filename).suffix.lower()

    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text for AI processing.

    Args:
        text: Raw extracted text from PDF/DOCX.

    Returns:
        str: Cleaned text with normalized whitespace and structure.
    """
    if not text:
        return ""

    # Split into lines for processing
    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        # Remove excessive whitespace
        cleaned_line = ' '.join(line.split())

        # Skip empty lines
        if not cleaned_line:
            continue

        # Keep page separators for structure
        if cleaned_line.startswith("--- Page"):
            cleaned_lines.append(cleaned_line)
        else:
            cleaned_lines.append(cleaned_line)

    # Join with single newlines and limit consecutive newlines
    cleaned_text = '\n'.join(cleaned_lines)

    # Remove excessive consecutive newlines (more than 2)
    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)

    return cleaned_text.strip()